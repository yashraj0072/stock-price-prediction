import yfinance as yf
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import train_test_split
from datetime import datetime
from docx import Document
from docx.shared import Inches
import os 
# Step 1: User input
ticker = input("Enter the stock ticker symbol (e.g., AAPL, TCS.NS): ").upper()

# Step 2: Download stock data
data = yf.download(ticker, period="1y", interval="1d")

# Step 3: Check if data is empty
if data.empty:
    print("❌ Failed to download data. Please check the ticker symbol.")
    exit()
# Step 3.1:
import os
folder_name = f"{ticker}_analysis_output"
os.makedirs(folder_name, exist_ok=True)


# Step 4: Calculate indicators
data['SMA_20'] = data['Close'].rolling(window=20).mean()
data['SMA_50'] = data['Close'].rolling(window=50).mean()
data['EMA_20'] = data['Close'].ewm(span=20, adjust=False).mean()
data['EMA_50'] = data['Close'].ewm(span=50, adjust=False).mean()

# Step 5: Fill missing values and drop early rows
data[['SMA_20', 'SMA_50', 'EMA_20', 'EMA_50']] = data[['SMA_20', 'SMA_50', 'EMA_20', 'EMA_50']].ffill()
data.dropna(inplace=True)

# Step 6: Determine trend
def determine_trend(row):
    try:
        if (
            float(row['Close']) > float(row['SMA_50']) and
            float(row['SMA_20']) > float(row['SMA_50']) and
            float(row['EMA_20']) > float(row['EMA_50'])
        ):
            return 'Bullish'
        elif (
            float(row['Close']) < float(row['SMA_50']) and
            float(row['SMA_20']) < float(row['SMA_50']) and
            float(row['EMA_20']) < float(row['EMA_50'])
        ):
            return 'Bearish'
        else:
            return 'Neutral'
    except:
        return 'Neutral'

# Apply the function to the data and create the 'Trend' column
data['Trend'] = data.apply(determine_trend, axis=1)

# Check if the 'Trend' column is added correctly
print(data.head())


# Step 7: Plot SMA
# Step 7: Plot SMA
sma_plot_path = os.path.join(folder_name, f"{ticker}_SMA_plot.png")

  # <-- Move this line here
plt.figure(figsize=(14, 6))
plt.plot(data['Close'], label='Close Price', color='black')
plt.plot(data['SMA_20'], label='SMA 20', color='blue')
plt.plot(data['SMA_50'], label='SMA 50', color='red')
plt.title('Simple Moving Averages')
plt.legend()
plt.grid(True)
plt.savefig(sma_plot_path)  # Save the plot as an image
plt.close()


# Step 8: Plot EMA
# Step 8: Plot EMA
ema_plot_path = os.path.join(folder_name, f"{ticker}_EMA_plot.png")
 # <-- Move this line here
plt.figure(figsize=(14, 6))
plt.plot(data['Close'], label='Close Price', color='black')
plt.plot(data['EMA_20'], label='EMA 20', color='green')
plt.plot(data['EMA_50'], label='EMA 50', color='orange')
plt.title('Exponential Moving Averages')
plt.legend()
plt.grid(True)
plt.savefig(ema_plot_path)  # Save the plot as an image
plt.close()
plt.show()

# Step 9: Print latest trend
latest_trend = data['Trend'].iloc[-1]
print(f"\n📈 Current Market Trend based on SMA and EMA: **{latest_trend}**")
# Step 9.5: Save data to CSV file
csv_path = os.path.join(folder_name, f"{ticker}_stock_data.csv")
data.to_csv(csv_path, index=True)

# Step 10: Regression target column
data['Target'] = (data['Open'] + data['Close']) / 2
data['Date_ordinal'] = pd.to_datetime(data.index).map(datetime.toordinal)

# Step 11: Linear Regression
X = data[['Date_ordinal']]
y = data['Target']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, shuffle=False)

model = LinearRegression()
model.fit(X_train, y_train)
# step 11.5
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

# Step 1: Predict on test data
y_pred = model.predict(X_test)

# Step 2: Calculate metrics
mae = mean_absolute_error(y_test, y_pred)
mse = mean_squared_error(y_test, y_pred)
r2 = r2_score(y_test, y_pred)

# Step 3: Display metrics
print(f"Model Accuracy Metrics:")
print(f"Mean Absolute Error (MAE): {mae:.2f}")
print(f"Mean Squared Error (MSE): {mse:.2f}")
print(f"R-squared (R²): {r2:.2f}")

# Step 12: Predict and plot
data['Predicted_Price'] = model.predict(X)

plt.figure(figsize=(14, 6))
plt.plot(data.index, data['Target'], label='Actual Avg Price', color='blue')
plt.plot(data.index, data['Predicted_Price'], label='Predicted Price', color='red')
plt.title('Actual vs Predicted Stock Prices')
plt.xlabel('Date')
plt.ylabel('Price')
plt.legend()
plt.grid(True)
plt.show()

# Step 12.5: Save data to Word file
doc_path = os.path.join(folder_name, f"{ticker}_report.docx")
doc = Document()
doc.add_heading(f'{ticker} Stock Report', 0)

# Add the content (Trend)
doc.add_paragraph(f"📅 Stock Data for {ticker}")
doc.add_paragraph(f"📈 Current Trend: {latest_trend}")

# Add graphs to Word
doc.add_picture(sma_plot_path, width=Inches(6))  # Add SMA plot
doc.add_picture(ema_plot_path, width=Inches(6))  # Add EMA plot

# Save the Word document
doc.save(doc_path)

# Step 13: Forecast future price
# Step 13: Forecast future price and trend
user_date = input("Enter a future date to forecast price (YYYY-MM-DD): ")
try:
    future_date = datetime.strptime(user_date, "%Y-%m-%d")
    future_ordinal = future_date.toordinal()
    forecast = model.predict([[future_ordinal]])

    # Calculate trend for the forecasted date
    future_data = pd.DataFrame({'Date_ordinal': [future_ordinal]})
    future_data['SMA_20'] = data['SMA_20'].iloc[-1]  # Last known SMA_20
    future_data['SMA_50'] = data['SMA_50'].iloc[-1]  # Last known SMA_50
    future_data['EMA_20'] = data['EMA_20'].iloc[-1]  # Last known EMA_20
    future_data['EMA_50'] = data['EMA_50'].iloc[-1]  # Last known EMA_50
    future_data['Close'] = forecast[0]  # Predicted close price

    # Trend calculation for future data
    future_trend = determine_trend(future_data.iloc[0])

    print(f"\n📅 Forecasted price for {user_date} is: ₹{forecast[0]:.2f}")
    print(f"📉 The predicted trend for {user_date} is: {future_trend}")

except Exception as e:
    print("❌ Invalid date format. Please use YYYY-MM-DD.")

